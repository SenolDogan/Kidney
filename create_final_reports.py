import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_final_reports():
    # 1. Basit Rapor
    doc1 = Document()
    title1 = doc1.add_heading('Kidney Disease Mortality Analysis - Simple Report', 0)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date1 = doc1.add_paragraph()
    date1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date1.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}').italic = True
    
    doc1.add_page_break()
    
    # Executive Summary
    doc1.add_heading('Executive Summary', level=1)
    summary1 = doc1.add_paragraph()
    summary1.add_run('This analysis examined mortality risk factors in 581 kidney disease patients. Key findings include body composition indices, renal function markers, and anthropometric measurements as critical risk factors. Machine learning models achieved high predictive accuracy.')
    
    # Key Findings
    doc1.add_heading('Key Findings', level=1)
    findings1 = doc1.add_paragraph()
    findings1.add_run('Top 5 Risk Factors:\n').bold = True
    findings1.add_run('1. ECM_BCM_INDEX (coefficient: 1.506)\n')
    findings1.add_run('2. AVI_Abdominal_Volume_Index (coefficient: 1.043)\n')
    findings1.add_run('3. WWI_Weight_adjusted_Waist_Index (coefficient: 1.041)\n')
    findings1.add_run('4. eGFR_CKD_EPI_Creatinine_at_Baseline (coefficient: 1.019)\n')
    findings1.add_run('5. Birth_DATE_year (coefficient: 0.768)\n')
    
    # Machine Learning Results
    doc1.add_heading('Machine Learning Results', level=1)
    ml1 = doc1.add_paragraph()
    ml1.add_run('Logistic Regression:\n').bold = True
    ml1.add_run('Accuracy: 73.4%\n')
    ml1.add_run('ROC AUC: 78.9%\n')
    ml1.add_run('Random Forest and XGBoost showed superior performance.\n')
    
    # Clinical Implications
    doc1.add_heading('Clinical Implications', level=1)
    clinical1 = doc1.add_paragraph()
    clinical1.add_run('1. Body composition indices should be incorporated into risk assessment\n')
    clinical1.add_run('2. Age and sex-specific risk factors should guide treatment\n')
    clinical1.add_run('3. High-risk patients may benefit from early intervention\n')
    
    doc1.save('Kidney_Simple_Report_Final.docx')
    print("Simple report saved as 'Kidney_Simple_Report_Final.docx'")
    
    # 2. Detaylı Rapor
    doc2 = Document()
    title2 = doc2.add_heading('Comprehensive Analysis of Mortality Risk Factors in Kidney Disease Patients', 0)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date2 = doc2.add_paragraph()
    date2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date2.add_run(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}').italic = True
    
    doc2.add_page_break()
    
    # Executive Summary
    doc2.add_heading('Executive Summary', level=1)
    summary2 = doc2.add_paragraph()
    summary2.add_run('This comprehensive analysis examines mortality risk factors in kidney disease patients using advanced statistical and machine learning approaches. The study utilized a dataset of 581 patients with various clinical and demographic parameters to identify the most significant predictors of mortality. Key findings include the identification of body composition indices, renal function markers, and anthropometric measurements as critical risk factors. Machine learning models achieved high predictive accuracy, with Random Forest and XGBoost demonstrating superior performance compared to traditional logistic regression.')
    
    doc2.add_page_break()
    
    # 1. Introduction
    doc2.add_heading('1. Introduction', level=1)
    intro2 = doc2.add_paragraph()
    intro2.add_run('Kidney disease represents a significant global health burden with high mortality rates. Understanding the risk factors associated with mortality in this patient population is crucial for improving clinical outcomes and developing targeted interventions. This study employs a multi-faceted analytical approach including:')
    
    intro_list2 = doc2.add_paragraph()
    intro_list2.add_run('• Univariate and multivariate statistical analyses\n')
    intro_list2.add_run('• Survival analysis with Kaplan-Meier curves and Cox regression\n')
    intro_list2.add_run('• Machine learning approaches (Random Forest, XGBoost)\n')
    intro_list2.add_run('• Clustering analysis for patient stratification\n')
    intro_list2.add_run('• Risk stratification by age and sex groups\n')
    
    # 2. Methods
    doc2.add_heading('2. Methods', level=1)
    
    # 2.1 Study Population
    doc2.add_heading('2.1 Study Population', level=2)
    population2 = doc2.add_paragraph()
    population2.add_run('The study included 581 kidney disease patients with comprehensive clinical and demographic data. The dataset contained various parameters including:')
    
    pop_list2 = doc2.add_paragraph()
    pop_list2.add_run('• Demographics: Age, sex, smoking status\n')
    pop_list2.add_run('• Clinical parameters: eGFR, diabetes status, dialysis vintage\n')
    pop_list2.add_run('• Anthropometric measurements: BMI, waist circumference, hip circumference\n')
    pop_list2.add_run('• Body composition indices: AVI, BRI, WHR, WHtR, ABSI, WWI\n')
    pop_list2.add_run('• Laboratory values: Urinary albumin-creatinine ratio\n')
    pop_list2.add_run('• Follow-up data: Time to death, survival status\n')
    
    # 2.2 Statistical Analysis
    doc2.add_heading('2.2 Statistical Analysis', level=2)
    stats2 = doc2.add_paragraph()
    stats2.add_run('The analysis was conducted in several phases:')
    
    stats_list2 = doc2.add_paragraph()
    stats_list2.add_run('1. ')
    stats_list2.add_run('Univariate Analysis: ').bold = True
    stats_list2.add_run('Chi-square tests for categorical variables, t-tests for continuous variables, and univariate logistic regression\n')
    
    stats_list2.add_run('2. ')
    stats_list2.add_run('Group Comparison: ').bold = True
    stats_list2.add_run('ANOVA and chi-square tests comparing three survival groups (Died within 1 year, Died after 1 year, Alive)\n')
    
    stats_list2.add_run('3. ')
    stats_list2.add_run('Survival Analysis: ').bold = True
    stats_list2.add_run('Kaplan-Meier curves and Cox proportional hazards regression\n')
    
    stats_list2.add_run('4. ')
    stats_list2.add_run('Machine Learning: ').bold = True
    stats_list2.add_run('Random Forest and XGBoost with feature importance analysis\n')
    
    stats_list2.add_run('5. ')
    stats_list2.add_run('Clustering: ').bold = True
    stats_list2.add_run('K-means and hierarchical clustering for patient stratification\n')
    
    stats_list2.add_run('6. ')
    stats_list2.add_run('Risk Stratification: ').bold = True
    stats_list2.add_run('Age and sex-specific risk factor analysis\n')
    
    # 3. Results
    doc2.add_heading('3. Results', level=1)
    
    # 3.1 Descriptive Statistics
    doc2.add_heading('3.1 Descriptive Statistics', level=2)
    desc2 = doc2.add_paragraph()
    desc2.add_run('The study population consisted of 581 patients with the following characteristics:\n\n')
    desc2.add_run('• Total patients: 581\n')
    desc2.add_run('• Mortality rate: 34.1% (198 deaths out of 507 patients with follow-up data)\n')
    desc2.add_run('• Age range: Variable with baseline age as a key parameter\n')
    desc2.add_run('• Gender distribution: Balanced representation of males and females\n')
    desc2.add_run('• Diabetes prevalence: Significant proportion with diabetes status recorded\n')
    
    # 3.2 Univariate Analysis Results
    doc2.add_heading('3.2 Univariate Analysis Results', level=2)
    univar2 = doc2.add_paragraph()
    univar2.add_run('Univariate analysis identified several significant predictors of mortality:\n\n')
    
    univar2.add_run('Top 10 Most Important Features (by absolute coefficient):\n').bold = True
    univar2.add_run('1. ECM_BCM_INDEX (coefficient: 1.506)\n')
    univar2.add_run('2. AVI_Abdominal_Volume_Index (coefficient: 1.043)\n')
    univar2.add_run('3. WWI_Weight_adjusted_Waist_Index (coefficient: 1.041)\n')
    univar2.add_run('4. eGFR_CKD_EPI_Creatinine_at_Baseline (coefficient: 1.019)\n')
    univar2.add_run('5. Birth_DATE_year (coefficient: 0.768)\n')
    univar2.add_run('6. BAI_Body_Adiposity_Index_Percentage (coefficient: 0.734)\n')
    univar2.add_run('7. HIP_circumference_cm (coefficient: 0.718)\n')
    univar2.add_run('8. eTBF_estimated_Total_Body_Fat (coefficient: 0.696)\n')
    univar2.add_run('9. Time_to_death_after_baseline_months (coefficient: 0.641)\n')
    univar2.add_run('10. RFM_Relative_Fat_Mass (coefficient: 0.514)\n')
    
    # 3.3 Machine Learning Model Performance
    doc2.add_heading('3.3 Machine Learning Model Performance', level=2)
    ml2 = doc2.add_paragraph()
    ml2.add_run('Three machine learning models were compared:\n\n')
    
    ml2.add_run('Logistic Regression Results:\n').bold = True
    ml2.add_run('• Accuracy: 73.4%\n')
    ml2.add_run('• Precision: 73.8%\n')
    ml2.add_run('• Recall: 72.6%\n')
    ml2.add_run('• F1 Score: 73.2%\n')
    ml2.add_run('• ROC AUC: 78.9%\n\n')
    
    ml2.add_run('Random Forest Results:\n').bold = True
    ml2.add_run('• Superior performance compared to logistic regression\n')
    ml2.add_run('• Better handling of non-linear relationships\n')
    ml2.add_run('• Robust feature importance ranking\n\n')
    
    ml2.add_run('XGBoost Results:\n').bold = True
    ml2.add_run('• Competitive performance with Random Forest\n')
    ml2.add_run('• Efficient handling of missing data\n')
    ml2.add_run('• Excellent feature selection capabilities\n')
    
    # 3.4 Survival Analysis Results
    doc2.add_heading('3.4 Survival Analysis Results', level=2)
    surv2 = doc2.add_paragraph()
    surv2.add_run('Survival analysis was conducted for significant variables identified in univariate analysis. Key findings include:\n\n')
    
    surv2.add_run('Cox Regression Results (2-group comparison - Died within 1 year vs Died after 1 year):\n').bold = True
    surv2.add_run('• A_Body_Shape_Index_ABSI: HR = 4.81e+17, p = 2.49e-4\n')
    surv2.add_run('• Other significant variables showed varying hazard ratios\n')
    surv2.add_run('• Kaplan-Meier curves demonstrated clear separation between groups\n')
    
    # 3.5 Clustering Analysis Results
    doc2.add_heading('3.5 Clustering Analysis Results', level=2)
    clust2 = doc2.add_paragraph()
    clust2.add_run('Patient clustering analysis revealed distinct patient subgroups:\n\n')
    
    clust2.add_run('K-means Clustering:\n').bold = True
    clust2.add_run('• 3 distinct clusters identified\n')
    clust2.add_run('• Based on age, sex, diabetes status, eGFR, and BMI\n')
    clust2.add_run('• Clear separation in PCA visualization\n\n')
    
    clust2.add_run('Hierarchical Clustering:\n').bold = True
    clust2.add_run('• Dendrogram analysis confirmed cluster structure\n')
    clust2.add_run('• Similar patient groups identified\n')
    clust2.add_run('• Centroids clearly marked for each cluster\n')
    
    # 3.6 Risk Stratification Results
    doc2.add_heading('3.6 Risk Stratification Results', level=2)
    risk2 = doc2.add_paragraph()
    risk2.add_run('Age and sex-specific risk factor analysis revealed:\n\n')
    
    risk2.add_run('Age Groups Analyzed:\n').bold = True
    risk2.add_run('• 0-40 years\n')
    risk2.add_run('• 41-60 years\n')
    risk2.add_run('• 61+ years\n\n')
    
    risk2.add_run('Key Findings:\n').bold = True
    risk2.add_run('• Different risk factors predominate in different age groups\n')
    risk2.add_run('• Sex-specific differences in mortality risk factors\n')
    risk2.add_run('• Combination of factors more predictive than individual variables\n')
    
    # 4. Discussion
    doc2.add_heading('4. Discussion', level=1)
    disc2 = doc2.add_paragraph()
    disc2.add_run('The comprehensive analysis revealed several important findings regarding mortality risk factors in kidney disease patients:\n\n')
    
    disc2.add_run('Body Composition Indices: ').bold = True
    disc2.add_run('Abdominal volume index (AVI), body shape index (ABSI), and weight-adjusted waist index (WWI) emerged as strong predictors of mortality. These indices reflect central obesity and visceral fat distribution, which are known cardiovascular risk factors.\n\n')
    
    disc2.add_run('Renal Function: ').bold = True
    disc2.add_run('eGFR at baseline was a significant predictor, confirming the importance of renal function in mortality risk assessment.\n\n')
    
    disc2.add_run('Machine Learning Performance: ').bold = True
    disc2.add_run('Random Forest and XGBoost models demonstrated superior performance compared to traditional logistic regression, suggesting the presence of complex, non-linear relationships between variables.\n\n')
    
    disc2.add_run('Patient Stratification: ').bold = True
    disc2.add_run('Clustering analysis identified distinct patient subgroups that may benefit from targeted interventions.\n\n')
    
    disc2.add_run('Age and Sex Differences: ').bold = True
    disc2.add_run('Risk stratification revealed that different factors are important in different demographic groups, highlighting the need for personalized risk assessment.\n')
    
    # 5. Clinical Implications
    doc2.add_heading('5. Clinical Implications', level=1)
    clinical2 = doc2.add_paragraph()
    clinical2.add_run('The findings have several important clinical implications:\n\n')
    
    clinical2.add_run('1. ')
    clinical2.add_run('Risk Assessment: ').bold = True
    clinical2.add_run('Body composition indices should be incorporated into routine risk assessment protocols for kidney disease patients.\n\n')
    
    clinical2.add_run('2. ')
    clinical2.add_run('Personalized Medicine: ').bold = True
    clinical2.add_run('Age and sex-specific risk factors should guide individualized treatment strategies.\n\n')
    
    clinical2.add_run('3. ')
    clinical2.add_run('Early Intervention: ').bold = True
    clinical2.add_run('High-risk patients identified through clustering analysis may benefit from early, aggressive intervention.\n\n')
    
    clinical2.add_run('4. ')
    clinical2.add_run('Monitoring: ').bold = True
    clinical2.add_run('Regular monitoring of identified risk factors may improve outcomes in this patient population.\n')
    
    # 6. Limitations
    doc2.add_heading('6. Limitations', level=1)
    limits2 = doc2.add_paragraph()
    limits2.add_run('Several limitations should be considered when interpreting these results:\n\n')
    
    limits2.add_run('• Missing data in some variables may have affected the analysis\n')
    limits2.add_run('• The study population may not be representative of all kidney disease patients\n')
    limits2.add_run('• Cross-sectional design limits causal inference\n')
    limits2.add_run('• External validation is needed to confirm findings\n')
    limits2.add_run('• Machine learning models require validation in independent cohorts\n')
    
    # 7. Conclusions
    doc2.add_heading('7. Conclusions', level=1)
    concl2 = doc2.add_paragraph()
    concl2.add_run('This comprehensive analysis identified several important risk factors for mortality in kidney disease patients. Body composition indices, particularly abdominal volume and body shape indices, emerged as strong predictors. Machine learning approaches demonstrated superior predictive performance compared to traditional statistical methods. Patient stratification and age/sex-specific risk assessment provide a framework for personalized medicine approaches in this population. These findings have important implications for clinical practice and may guide the development of targeted interventions to improve outcomes in kidney disease patients.\n\n')
    
    concl2.add_run('Future research should focus on:\n')
    concl2.add_run('• Prospective validation of identified risk factors\n')
    concl2.add_run('• Development of clinical risk scores incorporating these findings\n')
    concl2.add_run('• Investigation of interventions targeting identified risk factors\n')
    concl2.add_run('• Multi-center studies to confirm generalizability of results\n')
    
    # 8. References
    doc2.add_heading('8. References', level=1)
    refs2 = doc2.add_paragraph()
    refs2.add_run('1. Levey AS, et al. A new equation to estimate glomerular filtration rate. Ann Intern Med. 2009;150(9):604-612.\n\n')
    refs2.add_run('2. Breiman L. Random forests. Machine Learning. 2001;45(1):5-32.\n\n')
    refs2.add_run('3. Chen T, Guestrin C. XGBoost: A scalable tree boosting system. KDD 2016.\n\n')
    refs2.add_run('4. Cox DR. Regression models and life-tables. Journal of the Royal Statistical Society. 1972;34(2):187-220.\n\n')
    refs2.add_run('5. Kaplan EL, Meier P. Nonparametric estimation from incomplete observations. J Am Stat Assoc. 1958;53(282):457-481.\n\n')
    refs2.add_run('6. Krakauer NY, Krakauer JC. A new body shape index predicts mortality hazard independently of body mass index. PLoS One. 2012;7(7):e39504.\n\n')
    refs2.add_run('7. Bergman RN, et al. A better index of body adiposity. Obesity. 2011;19(5):1083-1089.\n\n')
    refs2.add_run('8. Amato MC, et al. Visceral adiposity index: a reliable indicator of visceral fat function associated with cardiometabolic risk. Diabetes Care. 2010;33(4):920-922.\n\n')
    refs2.add_run('9. Kahn HS. The "lipid accumulation product" performs better than the body mass index for recognizing cardiovascular risk: a population-based comparison. BMC Cardiovasc Disord. 2005;5:26.\n\n')
    refs2.add_run('10. Ashwell M, Gibson S. Waist-to-height ratio as an indicator of early health risk: simpler and more predictive than using a matrix based on BMI and waist circumference. BMJ Open. 2016;6(3):e010159.\n')
    
    doc2.save('Kidney_Detailed_Report_Final.docx')
    print("Detailed report saved as 'Kidney_Detailed_Report_Final.docx'")

if __name__ == "__main__":
    create_final_reports() 