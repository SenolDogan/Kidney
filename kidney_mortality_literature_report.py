from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_hyperlink(paragraph, url, text, color='0000FF', underline=True):
    # This function adds a hyperlink to a paragraph
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

doc = Document()

doc.add_heading('Mortality Risk Factors in Chronic Kidney Disease: A Comprehensive Survival Analysis', 0)

# Background
p = doc.add_heading('Background', level=1)
doc.add_paragraph(
    'Chronic kidney disease (CKD) patients are at high risk of mortality. Identifying which clinical and anthropometric parameters most strongly predict death can help improve patient management.'
)

# Methods
p = doc.add_heading('Methods', level=1)
methods = doc.add_paragraph()
methods.add_run('• Data from 581 CKD patients were analyzed.\n')
methods.add_run('• Patients were grouped as: died within 1 year, died after 1 year, or alive.\n')
methods.add_run('• Statistical tests: univariate logistic regression, t-test, chi-square, ANOVA, Kaplan-Meier, and Cox regression.\n')
methods.add_run('• Survival curves and hazard ratios were calculated for each significant variable.')

# Key Results
p = doc.add_heading('Key Results', level=1)
results = doc.add_paragraph()
results.add_run('• Strongest predictors of mortality: Lower eGFR, higher ABSI, WWI, ConI, BRI, older age, higher body fat, higher albuminuria, diabetes.\n')
results.add_run('• Novelty: Anthropometric indices (ABSI, WWI, etc.) are rarely reported in the literature as mortality predictors in CKD.\n')
results.add_run('• Survival group differences: Patients who died within 1 year had the worst risk profiles.\n')
results.add_run('• All results and plots are available on GitHub: https://github.com/SenolDogan/Kidney')

# Comparison with PubMed Literature
p = doc.add_heading('Comparison with PubMed Literature', level=1)

# A. Well-Established Risk Factors
p = doc.add_heading('A. Well-Established Risk Factors (Confirmed by Literature)', level=2)
well_established = doc.add_paragraph()
well_established.add_run('• eGFR: Lower eGFR is a key predictor (')
add_hyperlink(well_established, 'https://pubmed.ncbi.nlm.nih.gov/37591229/', 'Chen et al., 2023')
well_established.add_run(', ')
add_hyperlink(well_established, 'https://pubmed.ncbi.nlm.nih.gov/39334214/', 'Peng et al., 2024')
well_established.add_run(').\n')
well_established.add_run('• Age: Older age increases mortality risk (')
add_hyperlink(well_established, 'https://pubmed.ncbi.nlm.nih.gov/40259614/', 'Li et al., 2025')
well_established.add_run(').\n')
well_established.add_run('• Body composition: Obesity and body composition are linked to CKD outcomes, but indices like ABSI and WWI are less commonly reported.\n')
well_established.add_run('• Diabetes and blood glucose: Diabetes is a well-known risk factor (')
add_hyperlink(well_established, 'https://pubmed.ncbi.nlm.nih.gov/37591229/', 'Chen et al., 2023')
well_established.add_run(', ')
add_hyperlink(well_established, 'https://pubmed.ncbi.nlm.nih.gov/39915833/', 'Cao et al., 2025')
well_established.add_run(').\n')
well_established.add_run('• Albuminuria: Standard marker for CKD progression and mortality (')
add_hyperlink(well_established, 'https://pubmed.ncbi.nlm.nih.gov/39334214/', 'Peng et al., 2024')
well_established.add_run(').')

# B. More Original/Novel Findings
p = doc.add_heading('B. More Original/Novel Findings in Your Analysis', level=2)
novel = doc.add_paragraph()
novel.add_run('• ABSI, WWI, BRI, ConI: These indices are not commonly reported in large CKD mortality studies. Their strong association with mortality in your analysis is a novel contribution.\n')
novel.add_run('• Detailed survival grouping: Splitting deceased patients into "died within 1 year" and "died after 1 year" is more granular than most published studies.\n')
novel.add_run('• Comprehensive multi-method statistical approach.')

# C. Other Factors in Literature
p = doc.add_heading('C. Other Factors in Literature (Not Directly in Your Analysis)', level=2)
other = doc.add_paragraph()
other.add_run('• Lifestyle (physical activity, diet, smoking): Important in literature (')
add_hyperlink(other, 'https://pubmed.ncbi.nlm.nih.gov/37591229/', 'Chen et al., 2023')
other.add_run(', ')
add_hyperlink(other, 'https://pubmed.ncbi.nlm.nih.gov/39187191/', 'Tsai et al., 2024')
other.add_run('), but not detailed in your dataset.\n')
other.add_run('• Comorbidities (periodontitis, NAFLD, cognitive impairment): Recent studies show associations, but not directly analyzed in your dataset.')

# Conclusion
p = doc.add_heading('Conclusion', level=1)
doc.add_paragraph('Your analysis confirms known risk factors for mortality in CKD and introduces novel anthropometric indices as potential predictors. These findings may help refine risk stratification and guide future research.')

# References
p = doc.add_heading('References', level=1)
refs = doc.add_paragraph()
refs.add_run('• Chen et al., 2023, Am J Nephrol: https://pubmed.ncbi.nlm.nih.gov/37591229/\n')
refs.add_run('• Peng et al., 2024, BMC Med: https://pubmed.ncbi.nlm.nih.gov/39334214/\n')
refs.add_run('• Li et al., 2025, Ren Fail: https://pubmed.ncbi.nlm.nih.gov/40259614/\n')
refs.add_run('• Cao et al., 2025, Cardiovasc Diabetol: https://pubmed.ncbi.nlm.nih.gov/39915833/\n')
refs.add_run('• Tsai et al., 2024, J Affect Disord: https://pubmed.ncbi.nlm.nih.gov/39187191/')

doc.save('kidney_mortality_literature_report.docx')
print('Word report saved as kidney_mortality_literature_report.docx') 