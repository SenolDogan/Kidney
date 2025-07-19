import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_simple_report():
    # Yeni Word dokümanı oluştur
    doc = Document()
    
    # Başlık
    title = doc.add_heading('Kidney Disease Mortality Risk Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tarih
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}').italic = True
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This comprehensive analysis examined mortality risk factors in 581 kidney disease patients using advanced statistical and machine learning approaches. Key findings include identification of body composition indices, renal function markers, and anthropometric measurements as critical risk factors. Machine learning models achieved high predictive accuracy, with Random Forest and XGBoost demonstrating superior performance.')
    
    doc.add_page_break()
    
    # Study Overview
    doc.add_heading('1. Study Overview', level=1)
    overview = doc.add_paragraph()
    overview.add_run('Study Population: 581 kidney disease patients\n')
    overview.add_run('Mortality Rate: 34.1% (198 deaths out of 507 patients with follow-up data)\n')
    overview.add_run('Analysis Methods: Univariate analysis, Survival analysis, Machine learning, Clustering\n')
    
    # Key Findings
    doc.add_heading('2. Key Findings', level=1)
    
    # Top Risk Factors
    doc.add_heading('2.1 Top 10 Risk Factors', level=2)
    factors = doc.add_paragraph()
    factors.add_run('1. ECM_BCM_INDEX (coefficient: 1.506)\n')
    factors.add_run('2. AVI_Abdominal_Volume_Index (coefficient: 1.043)\n')
    factors.add_run('3. WWI_Weight_adjusted_Waist_Index (coefficient: 1.041)\n')
    factors.add_run('4. eGFR_CKD_EPI_Creatinine_at_Baseline (coefficient: 1.019)\n')
    factors.add_run('5. Birth_DATE_year (coefficient: 0.768)\n')
    factors.add_run('6. BAI_Body_Adiposity_Index_Percentage (coefficient: 0.734)\n')
    factors.add_run('7. HIP_circumference_cm (coefficient: 0.718)\n')
    factors.add_run('8. eTBF_estimated_Total_Body_Fat (coefficient: 0.696)\n')
    factors.add_run('9. Time_to_death_after_baseline_months (coefficient: 0.641)\n')
    factors.add_run('10. RFM_Relative_Fat_Mass (coefficient: 0.514)\n')
    
    # Machine Learning Results
    doc.add_heading('2.2 Machine Learning Model Performance', level=2)
    ml = doc.add_paragraph()
    ml.add_run('Logistic Regression Results:\n').bold = True
    ml.add_run('Accuracy: 73.4%\n')
    ml.add_run('Precision: 73.8%\n')
    ml.add_run('Recall: 72.6%\n')
    ml.add_run('F1 Score: 73.2%\n')
    ml.add_run('ROC AUC: 78.9%\n\n')
    
    ml.add_run('Random Forest and XGBoost showed superior performance compared to logistic regression.\n')
    
    # Survival Analysis
    doc.add_heading('2.3 Survival Analysis Results', level=2)
    surv = doc.add_paragraph()
    surv.add_run('Cox Regression Results (2-group comparison):\n').bold = True
    surv.add_run('A_Body_Shape_Index_ABSI: HR = 4.81e+17, p = 2.49e-4\n')
    surv.add_run('Kaplan-Meier curves demonstrated clear separation between groups.\n')
    
    # Clustering Results
    doc.add_heading('2.4 Clustering Analysis', level=2)
    clust = doc.add_paragraph()
    clust.add_run('K-means Clustering: 3 distinct clusters identified\n')
    clust.add_run('Based on: age, sex, diabetes status, eGFR, and BMI\n')
    clust.add_run('Clear separation in PCA visualization\n')
    
    # Clinical Implications
    doc.add_heading('3. Clinical Implications', level=1)
    clinical = doc.add_paragraph()
    clinical.add_run('1. Body composition indices should be incorporated into routine risk assessment\n')
    clinical.add_run('2. Age and sex-specific risk factors should guide treatment strategies\n')
    clinical.add_run('3. High-risk patients may benefit from early, aggressive intervention\n')
    clinical.add_run('4. Regular monitoring of identified risk factors may improve outcomes\n')
    
    # Conclusions
    doc.add_heading('4. Conclusions', level=1)
    concl = doc.add_paragraph()
    concl.add_run('This analysis identified several important risk factors for mortality in kidney disease patients. Body composition indices, particularly abdominal volume and body shape indices, emerged as strong predictors. Machine learning approaches demonstrated superior predictive performance. Patient stratification provides a framework for personalized medicine approaches.\n\n')
    
    concl.add_run('Future research should focus on:\n')
    concl.add_run('• Prospective validation of identified risk factors\n')
    concl.add_run('• Development of clinical risk scores\n')
    concl.add_run('• Investigation of targeted interventions\n')
    concl.add_run('• Multi-center validation studies\n')
    
    # Save the document
    doc.save('Kidney_Mortality_Simple_Report.docx')
    print("Simple report saved as 'Kidney_Mortality_Simple_Report.docx'")

if __name__ == "__main__":
    create_simple_report() 